# pip install unstructured[all-docs] pymupdf pdfplumber python-docx pandas
from __future__ import annotations

import os
import re
import json
import uuid
import argparse
from dataclasses import dataclass, asdict
from typing import Any, Optional, List, Dict, Iterable, Tuple

# -----------------------------
# Unified output data model
# -----------------------------

@dataclass
class Element:
    doc_id: str
    source_path: str
    parser: str                 # unstructured | pymupdf | pdfplumber | python-docx
    type: str                   # title | paragraph | list_item | table | page_break | other
    text: str
    page: Optional[int] = None  # 1-based when available
    bbox: Optional[Tuple[float, float, float, float]] = None  # (x0, y0, x1, y1) in PDF coords
    section_path: Optional[str] = None  # e.g. "1. Общие положения > 1.2 Термины"
    table: Optional[Dict[str, Any]] = None  # {rows: [[...]], markdown: "..."} optional
    meta: Optional[Dict[str, Any]] = None

    def citation(self) -> str:
        # Унифицированная ссылка-идентификатор на источник для RAG citations.
        # Для PDF: страница + bbox; для DOCX: без страницы, оставляем якорь.
        parts = [os.path.basename(self.source_path)]
        if self.page is not None:
            parts.append(f"p={self.page}")
        if self.bbox is not None:
            x0, y0, x1, y1 = self.bbox
            parts.append(f"bbox={round(x0,2)},{round(y0,2)},{round(x1,2)},{round(y1,2)}")
        return "#".join(parts)

# -----------------------------
# Helpers
# -----------------------------

def write_jsonl(path: str, items: Iterable[Element]) -> None:
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        for el in items:
            obj = asdict(el)
            obj["citation"] = el.citation()
            f.write(json.dumps(obj, ensure_ascii=False) + "\n")

def normalize_ws(s: str) -> str:
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()

def is_probable_list_item(text: str) -> bool:
    t = text.strip()
    return bool(re.match(r"^(\u2022|-|\*|\d+[\.\)]|[a-zA-Zа-яА-Я]\))\s+", t))

# -----------------------------
# PDF: Unstructured parser
# -----------------------------

def parse_pdf_unstructured(pdf_path: str, doc_id: Optional[str] = None) -> List[Element]:
    doc_id = doc_id or str(uuid.uuid4())

    from unstructured.partition.pdf import partition_pdf  # type: ignore

    # include_page_breaks=True + infer_table_structure=True + strategy="hi_res" —
    # практичная настройка для извлечения структуры/таблиц.
    elements = partition_pdf(
        filename=pdf_path,
        include_page_breaks=True,
        infer_table_structure=True,
        strategy="hi_res",
        ocr_languages="rus+eng",
    )

    out: List[Element] = []
    section_stack: List[str] = []

    for e in elements:
        etype = getattr(e, "category", None) or getattr(e, "type", None) or e.__class__.__name__
        text = normalize_ws(str(getattr(e, "text", "") or ""))

        meta = getattr(e, "metadata", None)
        page = getattr(meta, "page_number", None) if meta else None
        # bbox может называться по-разному; часто есть coordinates / bbox
        bbox = None
        if meta:
            coords = getattr(meta, "coordinates", None)
            if coords and hasattr(coords, "points") and coords.points:
                xs = [p[0] for p in coords.points]
                ys = [p[1] for p in coords.points]
                bbox = (min(xs), min(ys), max(xs), max(ys))

        # Rough mapping типов
        if str(etype).lower() in {"title"}:
            typ = "title"
            if text:
                # поддерживаем упрощенную иерархию: каждое Title считаем новым разделом верхнего уровня
                section_stack = [text]
        elif "list" in str(etype).lower():
            typ = "list_item"
        elif "table" in str(etype).lower():
            typ = "table"
        elif "pagebreak" in str(etype).lower() or "page_break" in str(etype).lower():
            typ = "page_break"
        else:
            typ = "paragraph"

        table_obj = None
        # В Unstructured таблица иногда хранится как text (HTML/строки).
        # Здесь сохраняем как есть; в downstream можно конвертировать HTML->markdown.
        if typ == "table":
            table_obj = {"raw": text}

        out.append(
            Element(
                doc_id=doc_id,
                source_path=pdf_path,
                parser="unstructured",
                type=typ,
                text=text,
                page=page,
                bbox=bbox,
                section_path=" > ".join(section_stack) if section_stack else None,
                table=table_obj,
                meta={"unstructured_type": str(etype)},
            )
        )

    return out

# -----------------------------
# PDF: PyMuPDF (fitz) parser
# -----------------------------

def parse_pdf_pymupdf(pdf_path: str, doc_id: Optional[str] = None) -> List[Element]:
    doc_id = doc_id or str(uuid.uuid4())
    import fitz  # PyMuPDF

    doc = fitz.open(pdf_path)
    out: List[Element] = []

    # Простая эвристика заголовков: размер шрифта существенно выше медианы на странице.
    for i in range(len(doc)):
        page = doc.load_page(i)
        d = page.get_text("dict")
        blocks = d.get("blocks", [])

        # собрать распределение шрифтов
        sizes = []
        for b in blocks:
            for line in b.get("lines", []):
                for span in line.get("spans", []):
                    if "size" in span:
                        sizes.append(float(span["size"]))
        median = sorted(sizes)[len(sizes)//2] if sizes else 0.0

        for b in blocks:
            if b.get("type") != 0:  # не текстовый блок
                continue
            bbox_block = tuple(b.get("bbox", (None, None, None, None)))
            lines = []
            max_span_size = 0.0
            for line in b.get("lines", []):
                line_text = "".join([s.get("text", "") for s in line.get("spans", [])])
                line_text = normalize_ws(line_text)
                if line_text:
                    lines.append(line_text)
                for s in line.get("spans", []):
                    if "size" in s:
                        max_span_size = max(max_span_size, float(s["size"]))

            text = normalize_ws("\n".join(lines))
            if not text:
                continue

            if median and max_span_size >= (median + 2.5):
                typ = "title"
            elif is_probable_list_item(text):
                typ = "list_item"
            else:
                typ = "paragraph"

            out.append(
                Element(
                    doc_id=doc_id,
                    source_path=pdf_path,
                    parser="pymupdf",
                    type=typ,
                    text=text,
                    page=i + 1,  # 1-based
                    bbox=bbox_block if all(v is not None for v in bbox_block) else None,
                    meta={"median_font_size": median, "max_span_size": max_span_size},
                )
            )

    return out

# -----------------------------
# PDF: pdfplumber parser (tables + text)
# -----------------------------

def parse_pdf_pdfplumber(pdf_path: str, doc_id: Optional[str] = None) -> List[Element]:
    doc_id = doc_id or str(uuid.uuid4())
    import pdfplumber

    out: List[Element] = []

    with pdfplumber.open(pdf_path) as pdf:
        for page_idx, page in enumerate(pdf.pages, start=1):
            # 1) Таблицы
            try:
                tables = page.find_tables()
            except Exception:
                tables = []

            for t_i, t in enumerate(tables):
                rows = t.extract()
                rows = rows or []
                out.append(
                    Element(
                        doc_id=doc_id,
                        source_path=pdf_path,
                        parser="pdfplumber",
                        type="table",
                        text="",
                        page=page_idx,
                        bbox=tuple(t.bbox) if getattr(t, "bbox", None) else None,
                        table={"rows": rows, "index_on_page": t_i},
                    )
                )

            # 2) Текст (можно исключать слова внутри таблиц, но оставим упрощенно)
            text = normalize_ws(page.extract_text() or "")
            if text:
                out.append(
                    Element(
                        doc_id=doc_id,
                        source_path=pdf_path,
                        parser="pdfplumber",
                        type="paragraph",
                        text=text,
                        page=page_idx,
                        bbox=None,
                    )
                )

    return out

# -----------------------------
# DOCX: python-docx parser
# -----------------------------

def parse_docx_python_docx(docx_path: str, doc_id: Optional[str] = None) -> List[Element]:
    doc_id = doc_id or str(uuid.uuid4())
    from docx import Document  # python-docx

    doc = Document(docx_path)
    out: List[Element] = []
    section_stack: List[str] = []

    # Параграфы
    for p_i, p in enumerate(doc.paragraphs):
        text = normalize_ws(p.text or "")
        if not text:
            continue

        style_name = (p.style.name or "").lower() if p.style else ""
        if "heading" in style_name:
            typ = "title"
            section_stack = [text]
        elif "list" in style_name or is_probable_list_item(text):
            typ = "list_item"
        else:
            typ = "paragraph"

        out.append(
            Element(
                doc_id=doc_id,
                source_path=docx_path,
                parser="python-docx",
                type=typ,
                text=text,
                page=None,   # у DOCX нет надежного page_number без рендеринга
                bbox=None,
                section_path=" > ".join(section_stack) if section_stack else None,
                meta={"paragraph_index": p_i, "style": p.style.name if p.style else None},
            )
        )

    # Таблицы
    for t_i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([normalize_ws(cell.text or "") for cell in row.cells])

        out.append(
            Element(
                doc_id=doc_id,
                source_path=docx_path,
                parser="python-docx",
                type="table",
                text="",
                table={"rows": rows, "table_index": t_i},
                meta={"section_path": " > ".join(section_stack) if section_stack else None},
            )
        )

    return out

# -----------------------------
# Orchestrator + CLI
# -----------------------------

def parse_document(path: str, prefer: str = "unstructured") -> Dict[str, List[Element]]:
    ext = os.path.splitext(path.lower())[1]
    results: Dict[str, List[Element]] = {}

    if ext == ".pdf":
        if prefer == "unstructured":
            results["unstructured"] = parse_pdf_unstructured(path)
            results["pymupdf"] = parse_pdf_pymupdf(path)
            results["pdfplumber"] = parse_pdf_pdfplumber(path)
        elif prefer == "pymupdf":
            results["pymupdf"] = parse_pdf_pymupdf(path)
            results["pdfplumber"] = parse_pdf_pdfplumber(path)
            results["unstructured"] = parse_pdf_unstructured(path)
        else:
            results["pdfplumber"] = parse_pdf_pdfplumber(path)
            results["pymupdf"] = parse_pdf_pymupdf(path)
            results["unstructured"] = parse_pdf_unstructured(path)

    elif ext == ".docx":
        results["python-docx"] = parse_docx_python_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return results

from pathlib import Path

def collect_inputs(input_dir: str, exts=(".pdf", ".docx")) -> list[str]:
    p = Path(input_dir)
    if not p.exists():
        raise FileNotFoundError(f"Input dir not found: {p.resolve()}")
    files = []
    for ext in exts:
        files.extend(p.rglob(f"*{ext}"))
        files.extend(p.rglob(f"*{ext.upper()}"))
    # сортировка стабильная, имена с русскими буквами ок
    return [str(f) for f in sorted(set(files))]

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input_dir", default="samples", help="Folder with PDF/DOCX files")
    ap.add_argument("--out_dir", default="out_parse", help="Output directory")
    ap.add_argument("--prefer", default="pymupdf", choices=["unstructured", "pymupdf", "pdfplumber"])
    args = ap.parse_args()

    inputs = collect_inputs(args.input_dir)
    if not inputs:
        raise RuntimeError(f"No PDF/DOCX found in: {args.input_dir}")

    for path in inputs:
        doc_name = Path(path).stem
        parsed = parse_document(path, prefer=args.prefer)

        base_dir = os.path.join(args.out_dir, doc_name)
        os.makedirs(base_dir, exist_ok=True)

        for parser_name, els in parsed.items():
            write_jsonl(os.path.join(base_dir, f"{parser_name}.jsonl"), els)

        preview_path = os.path.join(base_dir, "preview.md")
        with open(preview_path, "w", encoding="utf-8") as f:
            f.write(f"# Preview: {doc_name}\n\n")
            for parser_name, els in parsed.items():
                f.write(f"## {parser_name}\n\n")
                for el in els[:20]:
                    f.write(f"- [{el.type}] {el.citation()} :: {el.text[:200].replace('\\n',' ')}\n")
                f.write("\n")

if __name__ == "__main__":
    main()
